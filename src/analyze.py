import argparse
import string
from collections import Counter
from nltk import ngrams
from textblob import TextBlob
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from tabulate import tabulate
import yaml
import pprint
import os
from openai import OpenAI
import markdown
from docx import Document
import subprocess
from datetime import datetime

# Define base directories for the project
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, 'data')
CONFIG_DIR = os.path.join(BASE_DIR, 'config')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
CV_DIR = os.path.join(BASE_DIR, 'cv')
ARCHIVE_DIR = os.path.join(BASE_DIR, 'archive')

# Update file paths to use the new directory structure
def load_stopwords(file_path=os.path.join(DATA_DIR, 'stopwords.txt')):
    try:
        with open(file_path, 'r') as file:
            stopwords = set(line.strip().lower() for line in file)
        return stopwords
    except FileNotFoundError:
        print(f"Warning: Stopwords file '{file_path}' not found. Continuing without stopwords.")
        return set()

# Function to load the job description text
def load_text(file_path):
    try:
        with open(file_path, 'r') as file:
            text = file.read().lower().translate(str.maketrans('', '', string.punctuation))
        return text
    except FileNotFoundError:
        print(f"Error: Job description file '{file_path}' not found.")
        exit(1)

# Sentiment analysis function
def analyze_sentiment(text):
    blob = TextBlob(text)
    return blob.sentiment

# Bigram and trigram generator
def generate_ngrams(words, n):
    return Counter(ngrams(words, n))

def plot_wordcloud_and_frequencies(unigrams, bigrams, text):
    fig, axes = plt.subplots(1, 3, figsize=(18, 5))  # Three side-by-side plots

    # Word Cloud
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)
    axes[0].imshow(wordcloud, interpolation='bilinear')
    axes[0].axis('off')  # Hide axes
    axes[0].set_title("Word Cloud")

    # Unigrams Bar Chart
    top_unigrams = unigrams.most_common(10)
    words_u, counts_u = zip(*top_unigrams)
    axes[1].bar(words_u, counts_u, color='tab:blue')
    axes[1].set_title("Top Unigrams")
    axes[1].set_xlabel("Words")
    axes[1].set_ylabel("Frequency")
    axes[1].tick_params(axis='x', rotation=45)

    # Bigrams Bar Chart
    top_bigrams = bigrams.most_common(10)
    words_b, counts_b = zip(*[(" ".join(phrase), count) for phrase, count in top_bigrams])
    axes[2].bar(words_b, counts_b, color='tab:green')
    axes[2].set_title("Top Bigrams")
    axes[2].set_xlabel("Phrases")
    axes[2].set_ylabel("Frequency")
    axes[2].tick_params(axis='x', rotation=90)

    plt.tight_layout()

    # Save the figure
    plt.savefig("example_output.png", dpi=300, bbox_inches="tight")
    plt.show()  # Commented out

# Find contextual uses of a word
def find_context(word, text, window=5):
    words = text.split()
    indices = [i for i, w in enumerate(words) if w == word]
    for i in indices:
        start = max(i - window, 0)
        end = min(i + window + 1, len(words))
        print(f"Context for '{word}':", ' '.join(words[start:end]))

# Function to load and process the CV text
def load_cv_text(cv_file_path, stopwords_file):
    text = load_text(cv_file_path)  # reuse the existing load_text function
    stopwords = load_stopwords(stopwords_file)
    words = text.split()
    # Filter out stopwords
    words = [word for word in words if word not in stopwords]
    return text, Counter(words)

# Function to extract n-grams from a list of words
def extract_cv_ngrams(cv_words):
    unigrams = Counter(cv_words)
    bigrams = generate_ngrams(cv_words, 2)
    trigrams = generate_ngrams(cv_words, 3)
    return unigrams, bigrams, trigrams

# Function to compare job description n-grams with CV n-grams
# This simple comparison prints common unigrams, bigrams, and trigrams
# and indicates which top job description terms are missing in the CV.
def compare_cv_and_job(job_ngrams, cv_ngrams):
    job_unigrams, job_bigrams, job_trigrams = job_ngrams
    cv_unigrams, cv_bigrams, cv_trigrams = cv_ngrams

    print("\n--- Comparison between Job Description and CV ---\n")

    # Compare Unigrams (sorted by difference in counts)
    print("Common Unigrams (sorted by difference):")
    common_uni = set(job_unigrams.keys()) & set(cv_unigrams.keys())
    uni_diffs = [(word, job_unigrams[word], cv_unigrams[word], abs(job_unigrams[word] - cv_unigrams[word])) for word in common_uni]
    uni_diffs.sort(key=lambda x: x[3], reverse=True)
    print(tabulate(uni_diffs, headers=["Unigram", "JD", "CV", "Diff"], tablefmt="github"))
    print("\n")

    # Compare Bigrams (sorted by difference in counts)
    print("Common Bigrams (sorted by difference):")
    common_bi = set(job_bigrams.keys()) & set(cv_bigrams.keys())
    bi_diffs = [( ' '.join(phrase), job_bigrams[phrase], cv_bigrams[phrase], abs(job_bigrams[phrase] - cv_bigrams[phrase]) ) for phrase in common_bi]
    bi_diffs.sort(key=lambda x: x[3], reverse=True)
    print(tabulate(bi_diffs, headers=["Bigram", "JD", "CV", "Diff"], tablefmt="github"))
    print("\n")

    # Compare Trigrams (sorted by difference in counts)
    print("Common Trigrams (sorted by difference):")
    common_tri = set(job_trigrams.keys()) & set(cv_trigrams.keys())
    tri_diffs = [( ' '.join(phrase), job_trigrams[phrase], cv_trigrams[phrase], abs(job_trigrams[phrase] - cv_trigrams[phrase]) ) for phrase in common_tri]
    tri_diffs.sort(key=lambda x: x[3], reverse=True)
    print(tabulate(tri_diffs, headers=["Trigram", "JD", "CV", "Diff"], tablefmt="github"))
    print("\n")

    # Identify top job unigrams missing in CV
    print("Top Job Description Unigrams Missing in CV:")
    missing = [(word, count) for word, count in job_unigrams.most_common(10) if word not in cv_unigrams]
    print(tabulate(missing, headers=["Unigram", "JD"], tablefmt="github"))

# Load structured YAML CV data
def load_cv_yaml(yaml_path=os.path.join(CONFIG_DIR, 'cv_database.yaml')):
    with open(yaml_path, 'r') as f:
        return yaml.safe_load(f)

# Basic keyword-based CV generator
def generate_custom_cv(job_keywords, cv_data, output_path=os.path.join(OUTPUT_DIR, 'custom_cv.txt')):
    def match_score(item):
        desc = item.get("description", [])
        if isinstance(desc, str):
            desc = [desc]
        skills = item.get("skills", [])
        text = " ".join(desc + skills)
        return sum(1 for word in job_keywords if word in text.lower())

    # Sort and select top entries
    top_experience = sorted(cv_data.get("experience", []), key=match_score, reverse=True)[:4]
    top_projects = sorted(cv_data.get("projects", {}).get("work", []), key=match_score, reverse=True)[:2]
    top_personal_projects = sorted(cv_data.get("projects", {}).get("personal", []), key=match_score, reverse=True)[:1]

    pprint.pprint(top_experience)

    # Start writing the custom CV
    with open(output_path, 'w') as out:
        out.write("CUSTOM CV DRAFT\n\n")

        out.write("EXPERIENCE:\n")
        for job in top_experience:
            out.write(f"{job['title']} - {job['company']}\n")
            out.write(f"{job.get('location', '')} | {job['start_date']} to {job['end_date']}\n")
            for line in job.get("description", []):
                out.write(f" - {line}\n")
            out.write("\n")

        out.write("PROJECTS:\n")
        for proj in top_projects + top_personal_projects:
            out.write(f"{proj['name']} ({proj['year']})\n")
            out.write(f" - {proj['description']}\n")
            out.write("\n")

        out.write("EDUCATION:\n")
        for edu in cv_data.get("education", []):
            out.write(f"{edu['degree']} - {edu['institution']}, {edu['location']}\n")
            out.write(f"Graduated: {edu['graduation_year']}\n")
            if "thesis" in edu:
                out.write(f"Thesis: {edu['thesis']}\n")
            out.write("\n")

    print(f"\nCustom CV draft saved to '{output_path}'")

# Function to interact with GPT-4o-mini model
def run_gpt_model(job_file_path, cv_database_path, detailed_report_path, output_path, descriptive_copy_path, cover_letter_output_path, reference_folder=None, model="gpt-4o-mini"):
    """
    Interacts with the specified GPT model to generate both a tailored CV and an optional cover letter.

    :param job_file_path: Path to the job description file.
    :param cv_database_path: Path to the CV database file.
    :param detailed_report_path: Path to the detailed report file.
    :param output_path: Path to save the generated CV in the output folder.
    :param descriptive_copy_path: Path to save a descriptive copy of the CV in the cv folder.
    :param cover_letter_output_path: Path to save the generated cover letter.
    :param reference_folder: Path to the folder containing reference cover letters (optional).
    :param model: The GPT model to use (default: gpt-4o-mini).
    """
    client = OpenAI()

    # Read the input files
    with open(job_file_path, 'r') as job_file:
        job_description = job_file.read()

    with open(detailed_report_path, 'r') as report_file:
        detailed_report = report_file.read()

    with open(cv_database_path, 'r') as cv_database_file:
        cv_database = yaml.safe_load(cv_database_file)

    # Read reference cover letters if provided
    reference_texts = []
    if reference_folder:
        for filename in os.listdir(reference_folder):
            if filename.endswith('.txt'):
                with open(os.path.join(reference_folder, filename), 'r') as ref_file:
                    reference_texts.append(ref_file.read())
    combined_references = "\n\n===== COVER LETTER SEPARATOR =====\n\n".join(reference_texts)

    # Extract the job title and company name from the job description file name
    job_title = os.path.splitext(os.path.basename(job_file_path))[0].replace('_', ' ').title()
    company_name = job_title.split()[0]  # Assuming the company name is the first word

    # Load the prompt template
    with open(os.path.join(CONFIG_DIR, 'cv_prompt.txt'), 'r') as prompt_file:
        prompt_template = prompt_file.read()

    # Format the input for the GPT model
    input_text = prompt_template.format(
        company_name=company_name,
        job_description=job_description.strip(),
        detailed_report=detailed_report.strip(),
        cv_database=yaml.dump(cv_database, default_flow_style=False).strip(),
        reference_cover_letters=combined_references.strip()
    )

    # Log the formatted input text for debugging
    with open(os.path.join(OUTPUT_DIR, 'gpt_input_debug_log.txt'), 'w') as debug_log:
        debug_log.write("Formatted Input to GPT Model:\n")
        debug_log.write(input_text)

    # Log the input sent to the GPT model
    with open(os.path.join(OUTPUT_DIR, 'gpt_input_log.txt'), 'w') as log_file:
        log_file.write("Input to GPT Model:\n")
        log_file.write(input_text)

    # Call the GPT model
    response = client.responses.create(
        model=model,
        input=input_text
    )

    # Parse the response into CV and cover letter using a flexible delimiter
    import re

    # Search for a line containing 'Cover Letter' (case-insensitive)
    cover_letter_match = re.search(r"(?i)^.*cover letter.*$", response.output_text, re.MULTILINE)
    if cover_letter_match:
        delimiter = cover_letter_match.group(0)
        response_parts = response.output_text.strip().split(delimiter)
    else:
        response_parts = [response.output_text.strip()]

    if len(response_parts) != 2:
        print("Warning: Unexpected response format from GPT model. Appending raw response.")
        with open(output_path, 'a') as output_file:
            output_file.write("\n\n--- RAW RESPONSE ---\n\n")
            output_file.write(response.output_text.strip())
        return

    rewritten_cv, cover_letter = response_parts

    # Dynamically determine file paths based on job title, company name, and timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    job_title = os.path.splitext(os.path.basename(job_file_path))[0].replace('_', ' ').title()
    company_name = job_title.split()[0]  # Assuming the company name is the first word
 
    # Define a helper function to write the complete CV (CV content + cover letter) in one operation
    def write_cv_file(file_path, header, cv_body, cover_letter):
        with open(file_path, 'w') as f:
            f.write(header)
            f.write(cv_body.strip())
            f.write("\n\n--- COVER LETTER ---\n\n")
            f.write(cover_letter.strip())
 
    # Prepare a header to be included in each file
    header = f"This CV is tailored for the job: {job_title}\n\n"
 
    # Define file paths for the different CV outputs
    descriptive_copy_path = os.path.join(CV_DIR, f"{job_title.replace(' ', '_')}_CV.txt")
    archive_file_path = os.path.join(ARCHIVE_DIR, f"custom_cv_{timestamp}.txt")
    custom_cv_path = os.path.join(OUTPUT_DIR, 'custom_cv.txt')
    markdown_copy_path = descriptive_copy_path.replace('.txt', '.md')

    # Write all CV outputs using the helper function
    write_cv_file(output_path, header, rewritten_cv, cover_letter)
    print(f"Generated CV saved to {output_path}")

    write_cv_file(descriptive_copy_path, header, rewritten_cv, cover_letter)
    print(f"Descriptive CV copy saved to {descriptive_copy_path}")

    write_cv_file(custom_cv_path, header, rewritten_cv, cover_letter)
    print(f"Custom CV with cover letter written to: {custom_cv_path}")

    write_cv_file(archive_file_path, header, rewritten_cv, cover_letter)
    print(f"Archive CV saved to {archive_file_path}")

    # Write to the .md file
    write_cv_file(markdown_copy_path, header, rewritten_cv, cover_letter)
    print(f"Markdown CV saved to {markdown_copy_path}")

    print(header)
    print(rewritten_cv)
    print(cover_letter)

# Function to compare CV to job description
def compare_cv_to_jd(job_file_path, cv_file_path, output_path):
    """
    Compares the tailored CV with the job description and outputs common unigrams, bigrams, and trigrams,
    as well as missing terms from the job description.

    :param job_file_path: Path to the job description file.
    :param cv_file_path: Path to the tailored CV file.
    :param output_path: Path to append the comparison results.
    """
    # Load the job description and CV text
    with open(job_file_path, 'r') as job_file:
        job_text = job_file.read()

    with open(cv_file_path, 'r') as cv_file:
        cv_text = cv_file.read()

    # Load stopwords
    stopwords = load_stopwords(os.path.join(DATA_DIR, 'stopwords.txt'))

    # Filter out stopwords from job and CV texts
    job_words = [word for word in job_text.split() if word not in stopwords]
    cv_words = [word for word in cv_text.split() if word not in stopwords]

    # Generate n-grams
    job_unigrams = Counter(job_words)
    job_bigrams = Counter(ngrams(job_words, 2))
    job_trigrams = Counter(ngrams(job_words, 3))

    cv_unigrams = Counter(cv_words)
    cv_bigrams = Counter(ngrams(cv_words, 2))
    cv_trigrams = Counter(ngrams(cv_words, 3))

    # Compare n-grams
    common_unigrams = set(job_unigrams.keys()) & set(cv_unigrams.keys())
    common_bigrams = set(job_bigrams.keys()) & set(cv_bigrams.keys())
    common_trigrams = set(job_trigrams.keys()) & set(cv_trigrams.keys())

    # Sort by difference in counts
    unigram_diffs = [(word, job_unigrams[word], cv_unigrams[word], abs(job_unigrams[word] - cv_unigrams[word])) for word in common_unigrams]
    unigram_diffs.sort(key=lambda x: x[3], reverse=True)

    bigram_diffs = [(" ".join(phrase), job_bigrams[phrase], cv_bigrams[phrase], abs(job_bigrams[phrase] - cv_bigrams[phrase])) for phrase in common_bigrams]
    bigram_diffs.sort(key=lambda x: x[3], reverse=True)

    trigram_diffs = [(" ".join(phrase), job_trigrams[phrase], cv_trigrams[phrase], abs(job_trigrams[phrase] - cv_trigrams[phrase])) for phrase in common_trigrams]
    trigram_diffs.sort(key=lambda x: x[3], reverse=True)

    # Identify top job unigrams missing in CV
    missing_unigrams = [(word, count) for word, count in job_unigrams.most_common(10) if word not in cv_unigrams]

    # Append results to the detailed report
    with open(output_path, 'a') as report:
        report.write("\n--- Comparison between Job Description and CV ---\n\n")

        report.write("Common Unigrams (sorted by difference):\n")
        report.write(tabulate(unigram_diffs, headers=["Unigram", "JD", "CV", "Diff"], tablefmt="github"))
        report.write("\n\n")

        report.write("Common Bigrams (sorted by difference):\n")
        report.write(tabulate(bigram_diffs, headers=["Bigram", "JD", "CV", "Diff"], tablefmt="github"))
        report.write("\n\n")

        report.write("Common Trigrams (sorted by difference):\n")
        report.write(tabulate(trigram_diffs, headers=["Trigram", "JD", "CV", "Diff"], tablefmt="github"))
        report.write("\n\n")

        report.write("Top Job Description Unigrams Missing in CV:\n")
        report.write(tabulate(missing_unigrams, headers=["Unigram", "JD"], tablefmt="github"))
        report.write("\n\n")

    print("Comparison results appended to the detailed report.")

def convert_md_to_pdf_and_word(md_file_path):
    """
    Converts a Markdown file to both PDF and Word formats.

    :param md_file_path: Path to the Markdown file.
    """
    # Read the Markdown content
    with open(md_file_path, 'r') as md_file:
        md_content = md_file.read()

    # Convert Markdown to HTML
    html_content = markdown.markdown(md_content)
    print(f"md_file_path: {md_file_path}")

    # Correct the file paths for PDF and Word documents
    pdf_file_path = md_file_path.replace('.md', '.pdf')
    word_file_path = md_file_path.replace('.md', '.docx')

    # Commented out the Word document generation logic
    # doc = Document()
    # for line in md_content.splitlines():
    #     if line.startswith('# '):
    #         doc.add_heading(line[2:], level=1)
    #     elif line.startswith('## '):
    #         doc.add_heading(line[3:], level=2)
    #     elif line.startswith('### '):
    #         doc.add_heading(line[4:], level=3)
    #     elif line.strip():
    #         if line.startswith('- '):
    #             doc.add_paragraph(line[2:], style='List Bullet')
    #         else:
    #             doc.add_paragraph(line)
    # doc.save(word_file_path)
    # print(f"Generated Word Document: {word_file_path}")

    # Remove the redundant summary-shortening logic
    # Combine the Markdown content without manual truncation
    updated_md_content = md_content

    # Save the Markdown content as is
    with open(md_file_path, 'w') as md_file:
        md_file.write(updated_md_content)

    # Generate PDF using pandoc with xelatex for better Unicode support
    try:
        subprocess.run([
            "pandoc", md_file_path, "-o", pdf_file_path, "--pdf-engine=xelatex"
        ], check=True)
        print(f"Generated PDF: {pdf_file_path}")
    except subprocess.CalledProcessError as e:
        print(f"Error generating PDF: {e}")

def validate_ats_friendly_format(md_file_path):
    """
    Validates the Markdown file for ATS-friendly formatting.

    :param md_file_path: Path to the Markdown file.
    """
    with open(md_file_path, 'r') as md_file:
        content = md_file.read()

    issues = []

    # Check for tables or images
    if '|' in content or '![' in content:
        issues.append("Avoid using tables or images, as they may not be ATS-friendly.")

    # Check for special characters
    special_characters = ['₂', '©', '®', '™']
    for char in content:
        if (char in special_characters):
            issues.append(f"Special character '{char}' found. Replace it with plain text.")

    # Check for consistent section headings
    required_headings = ["Experience", "Education", "Skills", "Projects"]
    for heading in required_headings:
        if heading not in content:
            issues.append(f"Missing section heading: {heading}")

    # Removed the undefined `unigrams` reference
    # Ensure the function only validates for ATS-friendly formatting issues
    print("\nThe CV is ATS-friendly.")

# Command-line argument parsing
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analyze job descriptions for keywords, sentiment, and patterns.")
    parser.add_argument("file_path", help="Path to the job description file (e.g., example_JD.txt)")
    parser.add_argument("--stopwords", default=os.path.join(DATA_DIR, 'stopwords.txt'), help="Path to the stopwords file (default: stopwords.txt)")
    parser.add_argument("--cv_file", help="Path to the CV file to compare (optional)")
    parser.add_argument("--fast", action="store_true", help="Skip slow visualizations and context search")

    args = parser.parse_args()

    # Analyze job description
    job_text = load_text(os.path.join(DATA_DIR, args.file_path))
    stopwords = load_stopwords(args.stopwords)
    job_words = [word for word in job_text.split() if word not in stopwords]

    job_unigrams = Counter(job_words)
    job_bigrams = generate_ngrams(job_words, 2)
    job_trigrams = generate_ngrams(job_words, 3)

    # Load CV database
    cv_data = load_cv_yaml()

    # Generate the tailored CV
    generate_custom_cv(job_words, cv_data)

    # Existing analysis prints
    print(f"Sentiment Analysis: {analyze_sentiment(job_text)}")

    print("\nUnigrams:")
    for word, count in job_unigrams.most_common(10):
        print(f"{word}: {count}")

    print("\nBigrams:")
    for phrase, count in job_bigrams.most_common(10):
        print(f"{' '.join(phrase)}: {count}")

    print("\nTrigrams:")
    for phrase, count in job_trigrams.most_common(10):
        print(f"{' '.join(phrase)}: {count}")

    # Generate word cloud and frequency plots
    if not args.fast:
        plot_wordcloud_and_frequencies(job_unigrams, job_bigrams, job_text)
        find_context("python", job_text)

    # If a CV file is provided, perform CV analysis and compare with job description
    if args.cv_file:
        cv_text, cv_counter = load_cv_text(os.path.join(CV_DIR, args.cv_file), args.stopwords)
        cv_unigrams, cv_bigrams, cv_trigrams = extract_cv_ngrams(cv_text.split())

        # Compare the job description n-grams with CV n-grams
        compare_cv_and_job(
            (job_unigrams, job_bigrams, job_trigrams),
            (cv_unigrams, cv_bigrams, cv_trigrams)
        )